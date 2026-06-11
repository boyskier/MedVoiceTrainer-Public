import json
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_report(session: dict, output_path: str) -> str:
    """Generate a .docx report from session data and save to output_path."""
    doc = Document()

    # Title
    title = doc.add_heading("MedVoiceTrainer Session Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Session info table
    doc.add_heading("Session Information", level=2)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Table Grid"
    created_at = session.get("created_at", "")
    try:
        from datetime import timezone
        dt = datetime.fromisoformat(created_at)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        date_str = dt.astimezone().strftime("%Y-%m-%d %H:%M")
    except Exception:
        date_str = created_at

    duration = session.get("duration_seconds") or 0  # NULL for unfinalized sessions
    rows_data = [
        ("Date", date_str),
        ("Mode", session.get("mode", "").capitalize()),
        ("Case / Scenario", session.get("case_name", "")),
        ("Duration", f"{duration // 60}m {duration % 60}s"),
        ("Voice Backend", session.get("voice_backend", "").capitalize()),
        ("Eval Template", session.get("eval_template", "N/A")),
    ]
    for i, (label, value) in enumerate(rows_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)

    doc.add_paragraph()

    # Score summary table
    analysis = {}
    if session.get("raw_claude_response"):
        try:
            analysis = json.loads(session["raw_claude_response"])
        except Exception:
            pass

    scores = analysis.get("overall_scores", {})
    self_scores = {
        "grammar": session.get("self_grammar"),
        "medical_accuracy": session.get("self_medical_accuracy"),
        "clinical_reasoning": session.get("self_clinical_reasoning"),
        "professionalism": session.get("self_professionalism"),
        "communication_fluency": session.get("self_fluency"),
        "fluency": session.get("self_fluency"),
    }
    delta = analysis.get("self_assessment_delta", {})

    if scores:
        doc.add_heading("Score Summary", level=2)
        score_table = doc.add_table(rows=1 + len(scores), cols=4)
        score_table.style = "Table Grid"
        headers = score_table.rows[0].cells
        headers[0].text = "Metric"
        headers[1].text = "Claude Score (/10)"
        headers[2].text = "Self Score (/10)"
        headers[3].text = "Delta"
        def _fmt(value, spec=".1f"):
            try:
                return format(float(value), spec)
            except (TypeError, ValueError):
                return "N/A"

        for i, (metric, score) in enumerate(scores.items()):
            row = score_table.rows[i + 1]
            row.cells[0].text = metric.replace("_", " ").title()
            row.cells[1].text = _fmt(score)
            row.cells[2].text = _fmt(self_scores.get(metric))
            row.cells[3].text = _fmt(delta.get(metric), "+.1f")
        doc.add_paragraph()

    # Checklist
    checklist = analysis.get("checklist_results")
    if checklist:
        doc.add_heading("Checklist", level=2)
        cl_table = doc.add_table(rows=1 + len(checklist), cols=4)
        cl_table.style = "Table Grid"
        hrow = cl_table.rows[0].cells
        hrow[0].text = "Item"
        hrow[1].text = "Required"
        hrow[2].text = "Result"
        hrow[3].text = "Evidence"
        for i, item in enumerate(checklist):
            row = cl_table.rows[i + 1]
            row.cells[0].text = item.get("item", "")
            row.cells[1].text = "Yes" if item.get("required") else "No"
            passed = item.get("passed")
            row.cells[2].text = "✓ Pass" if passed else "✗ Fail"
            row.cells[3].text = item.get("evidence") or ""
        doc.add_paragraph()

    # SOAP Note
    soap = analysis.get("soap_note")
    if soap:
        doc.add_heading("SOAP Note", level=2)
        for section_key, section_label in [("subjective", "S — Subjective"), ("objective", "O — Objective"),
                                            ("assessment", "A — Assessment"), ("plan", "P — Plan")]:
            doc.add_heading(section_label, level=3)
            doc.add_paragraph(soap.get(section_key, ""))

        ref_soap_str = session.get("reference_soap")
        if not ref_soap_str:
            raw_case = session.get("raw_case_json")
            if raw_case:
                try:
                    case_data = json.loads(raw_case)
                    ref_soap_str = json.dumps(case_data.get("reference_soap", {}))
                except Exception:
                    pass

        if ref_soap_str:
            try:
                ref_soap = json.loads(ref_soap_str) if isinstance(ref_soap_str, str) else ref_soap_str
                doc.add_heading("Reference SOAP (Model Answer)", level=3)
                p = doc.add_paragraph()
                for section_key in ["subjective", "objective", "assessment", "plan"]:
                    run = p.add_run(f"{section_key.upper()}: {ref_soap.get(section_key, '')}\n")
                    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            except Exception:
                pass
        doc.add_paragraph()

    # Corrections
    corrections = analysis.get("corrections")
    if corrections:
        doc.add_heading("Language Corrections", level=2)
        for i, corr in enumerate(corrections, 1):
            p = doc.add_paragraph(f"{i}. [Turn {corr.get('turn_index', '?')}] ")
            run_orig = p.add_run(corr.get("original", ""))
            run_orig.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            p.add_run(" → ")
            run_corr = p.add_run(corr.get("corrected", ""))
            run_corr.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            doc.add_paragraph(f"   Explanation: {corr.get('explanation', '')}")
        doc.add_paragraph()

    # Summary
    summary = analysis.get("summary_feedback")
    if summary:
        doc.add_heading("Summary Feedback", level=2)
        doc.add_paragraph(summary)
        doc.add_paragraph()

    # Anki Cards preview
    anki_cards = analysis.get("anki_cards", [])
    if anki_cards:
        doc.add_heading("Anki Cards Preview (first 5)", level=2)
        anki_table = doc.add_table(rows=1 + min(5, len(anki_cards)), cols=2)
        anki_table.style = "Table Grid"
        hrow = anki_table.rows[0].cells
        hrow[0].text = "Front"
        hrow[1].text = "Back"
        for i, card in enumerate(anki_cards[:5]):
            row = anki_table.rows[i + 1]
            row.cells[0].text = card.get("front", "")
            row.cells[1].text = card.get("back", "")

    doc.save(output_path)
    return output_path


def suggest_filename(session: dict) -> str:
    created_at = session.get("created_at", "")
    try:
        from datetime import timezone
        dt = datetime.fromisoformat(created_at)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        date_str = dt.astimezone().strftime("%Y%m%d")
    except Exception:
        date_str = "unknown"
    mode = session.get("mode", "session")
    import re
    case_name = session.get("case_name", "case").replace(" ", "_")
    case_name = re.sub(r'[<>:"/\\|?*]', "-", case_name)
    return f"{date_str}_{mode}_{case_name}.docx"
